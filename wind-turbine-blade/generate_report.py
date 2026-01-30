"""
Generate progress report by filling the Word template.
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_report():
    doc = Document()

    # Title
    title = doc.add_heading('UNDERGRADUATE PROJECT PROGRESS REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cover info
    doc.add_paragraph()
    doc.add_paragraph('Project Title: Wind Turbine Blade Defect Detection Based on YOLOv11')
    doc.add_paragraph('Student ID: [Your Student ID]')
    doc.add_paragraph('Supervisor Name: [Your Supervisor Name]')
    doc.add_paragraph('Student Name: [Your Name]')
    doc.add_paragraph('Student Major: [Your Major]')
    doc.add_paragraph('Date Submitted: [Date]')

    doc.add_page_break()

    # Chapter 1
    doc.add_heading('Chapter 1. Introduction', level=1)

    doc.add_heading('1.1 Background of Study', level=2)
    doc.add_paragraph(
        'Wind energy has become one of the most important renewable energy sources worldwide. '
        'Wind turbines are critical infrastructure for clean energy generation, but their blades '
        'are constantly exposed to harsh environmental conditions including UV radiation, rain, ice, '
        'and debris impacts. These factors lead to various types of surface defects such as cracks, '
        'damage, dirt accumulation, and paint peeling. Early detection of these defects is crucial '
        'for preventing catastrophic failures and reducing maintenance costs.'
    )
    doc.add_paragraph(
        'Traditional blade inspection methods rely on manual visual inspection, which is time-consuming, '
        'labor-intensive, and potentially dangerous. With the rapid development of deep learning and '
        'computer vision technologies, automated defect detection systems have emerged as promising '
        'alternatives for efficient and accurate blade inspection.'
    )

    doc.add_heading('1.2 Project Aim', level=2)
    doc.add_paragraph(
        'The aim of this project is to develop an automated defect detection system for wind turbine '
        'blades using the YOLOv11 object detection model, capable of identifying and localizing four '
        'types of surface defects: cracks, damage, dirt, and peeled paint.'
    )

    doc.add_heading('1.3 Project Objectives', level=2)
    doc.add_paragraph('1. To prepare and preprocess the wind turbine blade defect dataset for training deep learning models.')
    doc.add_paragraph('2. To implement and train a YOLOv11-based object detection model for multi-class defect detection.')
    doc.add_paragraph('3. To optimize model hyperparameters and training strategies to improve detection accuracy.')
    doc.add_paragraph('4. To evaluate model performance using standard metrics including mAP, precision, and recall.')
    doc.add_paragraph('5. To analyze the detection performance for each defect category and identify areas for improvement.')

    doc.add_heading('1.4 Project Overview', level=2)
    doc.add_heading('1.4.1 Scope', level=3)
    doc.add_paragraph(
        'This study focuses on developing a deep learning-based automated inspection system for wind '
        'turbine blade surface defects. The significance of this research lies in its potential to '
        'improve inspection efficiency, reduce maintenance costs, and enhance the safety and reliability '
        'of wind energy infrastructure.'
    )

    doc.add_heading('1.4.2 Audience', level=3)
    doc.add_paragraph(
        'The findings of this project will benefit wind farm operators, maintenance engineers, and '
        'researchers working on automated inspection systems for renewable energy infrastructure.'
    )

    doc.add_page_break()

    # Chapter 2
    doc.add_heading('Chapter 2. Literature Review', level=1)

    doc.add_heading('2.1 Fundamental Theories', level=2)
    doc.add_paragraph(
        'Object Detection: Object detection is a computer vision task that involves identifying and '
        'localizing objects within images. Modern object detection methods are primarily based on deep '
        'convolutional neural networks (CNNs).'
    )
    doc.add_paragraph(
        'YOLO (You Only Look Once): YOLO is a real-time object detection algorithm that processes the '
        'entire image in a single forward pass, making it significantly faster than region-based methods. '
        'YOLOv11 is the latest iteration, featuring improved architecture and training strategies.'
    )

    doc.add_heading('2.2 Background Review', level=2)
    doc.add_heading('2.2.1 Traditional Methods', level=3)
    doc.add_paragraph(
        'Traditional defect detection methods for wind turbine blades include manual visual inspection '
        'using drones or climbing equipment, ultrasonic testing and thermographic inspection, and image '
        'processing techniques using edge detection and morphological operations. These methods have '
        'limitations in terms of efficiency, scalability, and ability to detect subtle defects.'
    )

    doc.add_heading('2.2.2 Deep Learning-based Methods', level=3)
    doc.add_paragraph(
        'Recent advances in deep learning have led to significant improvements in automated defect '
        'detection, including CNN-based classification methods, region-based detectors (R-CNN, Faster R-CNN), '
        'and single-shot detectors (YOLO, SSD) for real-time detection applications.'
    )

    doc.add_heading('2.3 Gaps in the Existing Literature', level=2)
    doc.add_paragraph(
        'While existing research has demonstrated the potential of deep learning for defect detection, '
        'several gaps remain: limited publicly available datasets specific to wind turbine blade defects, '
        'class imbalance issues affecting detection of less common defect types, and need for lightweight '
        'models suitable for edge deployment.'
    )

    doc.add_page_break()

    # Chapter 3
    doc.add_heading('Chapter 3. Technical Progress', level=1)

    doc.add_heading('3.1 Problem Statement', level=2)
    doc.add_paragraph(
        'The problem addressed in this project is the automated detection and localization of surface '
        'defects on wind turbine blades from visual imagery. The system must identify four defect '
        'categories (crack, damage, dirt, peeled paint) with high accuracy while maintaining practical '
        'inference speeds.'
    )

    doc.add_heading('3.2 Approach', level=2)
    doc.add_heading('3.2.1 Model Architecture', level=3)
    doc.add_paragraph(
        'YOLOv11s (small variant) was selected as the base model, featuring CSPDarknet backbone for '
        'feature extraction, PANet neck for multi-scale feature fusion, and decoupled detection head '
        'for improved accuracy.'
    )

    doc.add_heading('3.2.2 Dataset Description', level=3)

    # Dataset table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Class'
    hdr[1].text = 'Train'
    hdr[2].text = 'Val'
    hdr[3].text = 'Test'

    data = [
        ('crack', '949', '93', '127'),
        ('damage', '1,288', '173', '131'),
        ('dirt', '1,010', '118', '109'),
        ('peeled_paint', '964', '153', '124'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row[j].text = cell_data

    doc.add_paragraph()
    doc.add_paragraph('Table 3-1. Dataset class distribution')

    doc.add_heading('3.3 Technology', level=2)
    doc.add_paragraph('Hardware: NVIDIA RTX 4060 (8GB VRAM)')
    doc.add_paragraph('Software: Python 3.x, Ultralytics YOLOv11, PyTorch, OpenCV')

    doc.add_heading('3.4 Experimental Results', level=2)

    doc.add_heading('3.4.1 V1 Baseline Results', level=3)

    # V1 results table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Metric'
    table.rows[0].cells[1].text = 'Value'
    table.rows[1].cells[0].text = 'mAP50'
    table.rows[1].cells[1].text = '74.1%'
    table.rows[2].cells[0].text = 'mAP50-95'
    table.rows[2].cells[1].text = '49.2%'
    table.rows[3].cells[0].text = 'Precision'
    table.rows[3].cells[1].text = '76.2%'
    table.rows[4].cells[0].text = 'Recall'
    table.rows[4].cells[1].text = '69.3%'

    doc.add_paragraph()
    doc.add_paragraph('Table 3-2. V1 training results')

    doc.add_heading('3.4.2 V2 Optimized Results', level=3)
    doc.add_paragraph(
        'Based on V1 analysis, the following optimizations were applied: increased epochs to 150, '
        'extended patience to 20, added data augmentation (mixup=0.1, copy_paste=0.1), and adjusted '
        'loss weights.'
    )

    # Comparison table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'V1'
    hdr[2].text = 'V2'
    hdr[3].text = 'Change'

    data = [
        ('mAP50', '74.1%', '74.8%', '+0.7%'),
        ('mAP50-95', '49.2%', '49.5%', '+0.3%'),
        ('Precision', '76.2%', '80.2%', '+4.0%'),
        ('Recall', '69.3%', '68.5%', '-0.8%'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row[j].text = cell_data

    doc.add_paragraph()
    doc.add_paragraph('Table 3-3. V1 vs V2 comparison')

    doc.add_heading('3.4.3 Per-class Performance', level=3)

    # Per-class table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Class'
    hdr[1].text = 'V1'
    hdr[2].text = 'V2'
    hdr[3].text = 'Change'

    data = [
        ('crack', '80%', '79%', '-1%'),
        ('damage', '86%', '87%', '+1%'),
        ('dirt', '38%', '39%', '+1%'),
        ('peeled_paint', '85%', '90%', '+5%'),
    ]
    for i, row_data in enumerate(data):
        row = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row[j].text = cell_data

    doc.add_paragraph()
    doc.add_paragraph('Table 3-4. Per-class accuracy comparison')

    doc.add_heading('3.5 Analysis and Discussion', level=2)
    doc.add_paragraph(
        'Strengths: High detection accuracy for crack (79%), damage (87%), and peeled_paint (90%). '
        'Precision improved significantly from V1 to V2 (+4%). Model converged well without overfitting.'
    )
    doc.add_paragraph(
        'Challenges: Dirt class detection remains problematic (39% recall). This is likely due to '
        'visual similarity between dirt and background, potential annotation inconsistencies, and '
        'class imbalance in training data.'
    )

    doc.add_heading('3.6 Summary', level=2)
    doc.add_paragraph(
        'Two versions of the YOLOv11s model were trained and evaluated. V2 achieved improved overall '
        'performance with mAP50 of 74.8% and mAP50-95 of 49.5%. The model performs well on crack, damage, '
        'and peeled_paint detection but struggles with dirt class due to data-related challenges.'
    )

    doc.add_page_break()

    # Chapter 4
    doc.add_heading('Chapter 4. Project Management', level=1)

    doc.add_heading('4.1 Activities', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Objective'
    table.rows[0].cells[1].text = 'Tasks'
    table.rows[1].cells[0].text = 'Data Preparation'
    table.rows[1].cells[1].text = 'Dataset analysis, format conversion, train/val/test split'
    table.rows[2].cells[0].text = 'Model Development'
    table.rows[2].cells[1].text = 'YOLOv11 implementation, training script development'
    table.rows[3].cells[0].text = 'Experiment V1'
    table.rows[3].cells[1].text = 'Baseline training, evaluation, result analysis'
    table.rows[4].cells[0].text = 'Experiment V2'
    table.rows[4].cells[1].text = 'Hyperparameter optimization, augmentation tuning'
    table.rows[5].cells[0].text = 'Documentation'
    table.rows[5].cells[1].text = 'Progress report writing, result visualization'

    doc.add_paragraph()
    doc.add_paragraph('Table 4-1. Project activities')

    doc.add_heading('4.2 Schedule', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Week'
    table.rows[0].cells[1].text = 'Activity'
    table.rows[0].cells[2].text = 'Status'
    table.rows[1].cells[0].text = 'Week 1-2'
    table.rows[1].cells[1].text = 'Literature review and dataset preparation'
    table.rows[1].cells[2].text = 'Completed'
    table.rows[2].cells[0].text = 'Week 3-4'
    table.rows[2].cells[1].text = 'Model implementation and V1 training'
    table.rows[2].cells[2].text = 'Completed'
    table.rows[3].cells[0].text = 'Week 5-6'
    table.rows[3].cells[1].text = 'V1 analysis and V2 optimization'
    table.rows[3].cells[2].text = 'Completed'
    table.rows[4].cells[0].text = 'Week 7-8'
    table.rows[4].cells[1].text = 'Further optimization and report writing'
    table.rows[4].cells[2].text = 'In Progress'

    doc.add_paragraph()
    doc.add_paragraph('Table 4-2. Project schedule')

    doc.add_page_break()

    # Chapter 5
    doc.add_heading('Chapter 5. Professional Issues and Risk', level=1)

    doc.add_heading('5.1 Risk Analysis', level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Risk'
    table.rows[0].cells[1].text = 'Likelihood'
    table.rows[0].cells[2].text = 'Impact'
    table.rows[0].cells[3].text = 'Mitigation'
    table.rows[1].cells[0].text = 'GPU memory limitations'
    table.rows[1].cells[1].text = 'High'
    table.rows[1].cells[2].text = 'Medium'
    table.rows[1].cells[3].text = 'Reduced batch size, used smaller model'
    table.rows[2].cells[0].text = 'Model overfitting'
    table.rows[2].cells[1].text = 'Medium'
    table.rows[2].cells[2].text = 'High'
    table.rows[2].cells[3].text = 'Early stopping, data augmentation'
    table.rows[3].cells[0].text = 'Poor class performance'
    table.rows[3].cells[1].text = 'High'
    table.rows[3].cells[2].text = 'Medium'
    table.rows[3].cells[3].text = 'Hyperparameter tuning'
    table.rows[4].cells[0].text = 'Dataset quality issues'
    table.rows[4].cells[1].text = 'Medium'
    table.rows[4].cells[2].text = 'High'
    table.rows[4].cells[3].text = 'Data analysis, annotation review'

    doc.add_paragraph()
    doc.add_paragraph('Table 5-1. Risk analysis')

    doc.add_heading('5.2 Professional Issues', level=2)
    doc.add_paragraph(
        'Ethical Considerations: The dataset used is publicly available for research purposes. '
        'Model development follows responsible AI practices.'
    )
    doc.add_paragraph(
        'Environmental Impact: This project contributes to sustainable energy by improving wind '
        'turbine maintenance efficiency. Reduced manual inspection requirements lower carbon footprint.'
    )

    doc.add_page_break()

    # References
    doc.add_heading('References', level=1)
    doc.add_paragraph('[1] Redmon, J., et al. "You Only Look Once: Unified, Real-Time Object Detection." CVPR, 2016.')
    doc.add_paragraph('[2] Ultralytics. "YOLOv11 Documentation." https://docs.ultralytics.com/')
    doc.add_paragraph('[3] Wang, C.Y., et al. "CSPNet: A New Backbone that can Enhance Learning Capability of CNN." CVPR, 2020.')

    # Save
    doc.save('wind-turbine-blade/progress_report.docx')
    print('Report saved to wind-turbine-blade/progress_report.docx')


if __name__ == '__main__':
    create_report()
